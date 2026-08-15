import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_hr_clarification_doc(file_path):
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styling colors
    PRIMARY_COLOR = RGBColor(27, 67, 50)     # #1B4332 Deep Emerald Green
    SECONDARY_COLOR = RGBColor(45, 138, 78)  # #2D8A4E Emerald
    TEXT_DARK = RGBColor(30, 41, 59)        # #1E293B Dark Slate
    MUTED_TEXT = RGBColor(100, 116, 139)    # #64748B Slate Muted

    # Normal Style Configuration
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_DARK

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_sub = p_title.add_run("NOTRE DAME OF MARBEL UNIVERSITY • ACHIEVENEST PORTAL\n")
    r_sub.font.size = Pt(9.5)
    r_sub.font.bold = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    r_title = p_title.add_run("HR Admin Architecture, Workflow Logic & Clarification Questions")
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(20)
    r_meta = p_meta.add_run("Prepared for Stakeholder Alignment • August 2026")
    r_meta.font.size = Pt(9.5)
    r_meta.font.italic = True
    r_meta.font.color.rgb = MUTED_TEXT

    # Helper functions
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = SECONDARY_COLOR
        return h

    # =========================================================================
    # PART 1: CURRENT VERSION & WORKFLOW LOGIC
    # =========================================================================
    add_heading_1("PART 1: Current HR Admin Version & Workflow Logic")

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.add_run(
        "The HR Admin module in AchieveNest handles faculty accounts, academic leadership assignments, "
        "portfolio verification, and CHEd / PACUCOA accreditation compliance."
    )

    # 1.1 Core Modules
    add_heading_2("1.1 Core HR Admin Modules")
    
    modules = [
        ("1. HR Command Center: ", "High-level metrics (Total Faculty Accounts, CHEd/PACUCOA Readiness Score, Pending Verification Queue)."),
        ("2. Personnel Directory: ", "Master directory of faculty, staff, and secretaries with manual creation and CSV import."),
        ("3. Verification Queue: ", "2-tier audit queue (Faculty portfolios reviewed by Dept Sec; Dept Sec portfolios reviewed directly by HR Admin)."),
        ("4. CHEd & PACUCOA Ranking Board: ", "Automated ranking board computing points across Area A (Professional Development), Area B (Productivity), and Area C (Service)."),
        ("5. Audit Logs & Records: ", "Tracking system logging role assignments, promotions, password resets, and security seals.")
    ]

    for title, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_t = p.add_run(title)
        r_t.font.bold = True
        r_t.font.color.rgb = PRIMARY_COLOR
        p.add_run(desc)

    # 1.2 Administrative Role Assignment Matrix
    add_heading_2("1.2 Administrative Role Assignment Matrix")
    
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Role / Position", "Assigned By", "Scope & Responsibilities"]
    col_widths = [Inches(2.2), Inches(1.3), Inches(2.8)]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        shading = parse_xml(r'<w:shd {} w:fill="1B4332"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)

    row_data = [
        ("College Dean", "HR Admin", "Oversees an entire College (CEAC, CBA, CAS, CED)."),
        ("Program Director / Coordinator", "HR Admin", "Oversees a Degree Program (e.g. BS Computer Science)."),
        ("Department Secretary", "HR Admin", "Evaluates and endorses faculty portfolios under their department."),
        ("Student Organization Moderator", "OSAD Admin", "Serves as Faculty Advisor for a student club or organization.")
    ]

    for row_idx, data in enumerate(row_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths[col_idx]
            if row_idx % 2 == 0:
                shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                row_cells[col_idx]._tc.get_or_add_tcPr().append(shading)
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1.3 Sequential Workflow Order
    add_heading_2("1.3 Master Sequential Setup Workflow Order")

    flow_steps = [
        ("1. Create Academic Departments & Degree Programs", "HR / OSAD Admin"),
        ("2. Onboard Faculty & Personnel Accounts", "HR Admin"),
        ("3. Assign College Deans", "HR Admin"),
        ("4. Assign Program Directors / Coordinators", "HR Admin"),
        ("5. Assign Department Secretaries", "HR Admin"),
        ("6. Import / Create Student Accounts", "OSAD Admin"),
        ("7. Create Student Organizations & Clubs", "OSAD Admin"),
        ("8. Assign Organization Moderators", "OSAD Admin")
    ]

    for step_title, step_actor in flow_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_st = p.add_run(step_title)
        r_st.font.bold = True
        r_st.font.color.rgb = PRIMARY_COLOR
        p.add_run(f" ({step_actor})")

    # =========================================================================
    # PART 2: QUESTIONS FOR CLARIFICATION (UPDATED BASED ON USER FEEDBACK)
    # =========================================================================
    add_heading_1("PART 2: Questions for Clarification")

    categories = [
        ("Category A: Institutional Workflow & Setup Flow", [
            ("1. Setup Flow Order: ", "Does the 8-step setup order (Departments → Faculty → Deans → Coordinators → Dept Secs → Students → Orgs → Moderators) match NDMU's official annual administrative workflow?"),
            ("2. Assignment Sequence Flexibility: ", "Can HR Admin assign Deans, Coordinators, and Dept Secs in any order during setup, or must it strictly follow the step-by-step sequence?"),
            ("3. Mid-Year Department Additions: ", "If a new Department or Program is added mid-year, can OSAD Admin immediately map Student Organizations to it without re-running the full setup?")
        ]),

        ("Category B: Role Expiration & Term Limits", [
            ("4. Role Expiration Terms: ", "What is the exact expiration period or term limit for each role (College Dean, Program Coordinator, Department Secretary, Organization Moderator)?")
        ]),

        ("Category C: Portfolio Verification & Review Flow", [
            ("5. Dean Involvement in Review Flow: ", "Since regular faculty portfolios are reviewed by the Dept Sec, and Dept Sec portfolios are reviewed directly by HR, does the College Dean have any mandatory review step in this flow?"),
            ("6. Mandatory Dept Sec Review: ", "Is Dept Sec review strictly required for regular faculty before HR final audit, or can HR audit a faculty portfolio directly if needed?"),
            ("7. Returned Portfolios: ", "How many working days do faculty have to re-submit proof if their portfolio is returned for revision?"),
            ("8. Security Seals: ", "Should security seal codes (HR-SEAL-2026-XXXX) be attached to each individual proof item, or to the final annual ranking report?")
        ]),

        ("Category D: Account Management & Security", [
            ("9. Inactive Faculty Accounts: ", "When a faculty member goes on leave or resigns, what happens to their assigned administrative roles?"),
            ("10. Temporary Passwords: ", "Should temporary passwords generated during password resets expire after 24 hours?")
        ])
    ]

    for cat_title, q_list in categories:
        add_heading_2(cat_title)
        for q_title, q_body in q_list:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            r_qn = p.add_run(q_title)
            r_qn.font.bold = True
            r_qn.font.color.rgb = PRIMARY_COLOR
            p.add_run(q_body)

    # Save document with lock handling
    try:
        doc.save(file_path)
        print(f"Successfully updated DOCX at: {file_path}")
    except PermissionError:
        alt_path = file_path.replace(".docx", " (Simplified).docx")
        doc.save(alt_path)
        print(f"File open in another program. Saved to alternative path: {alt_path}")

if __name__ == '__main__':
    target = r"C:\Users\Admin\Downloads\AchieveNest  Docs\HR Follow up Questions for Clarification.docx"
    create_hr_clarification_doc(target)
